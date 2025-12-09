"""
Generate volunteer data templates for HOPE Tutoring
Creates empty templates and demo versions in CSV, DOCX, and PDF formats
"""

import csv
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Column headers based on HOPE Tutoring volunteer form
HEADERS = [
    "Last Name",
    "First Name", 
    "Email Address",
    "Phone Number",
    "Street Address with City & Zip Code",
    "Emergency Contact Name & Phone Number",
    "Are you at least 18 years old?",
    "Assigned Site/Session",
    "How did you hear about HOPE Tutoring?"
]

# Shortened headers for table display (more readable)
SHORT_HEADERS = [
    "Last Name",
    "First Name", 
    "Email Address",
    "Phone Number",
    "Address (City, Zip)",
    "Emergency Contact",
    "18+?",
    "Site/Session",
    "How did you hear about us?"
]

# Demo data - realistic sample volunteers (Batch 1)
DEMO_DATA = [
    ["Johnson", "Sarah", "sarah.johnson@email.com", "817-555-0101", "123 Oak St, Arlington, TX 76010", "Mike Johnson, 817-555-0102", "Yes", "Monday Evening - UTA", "Social Media"],
    ["Chen", "Michael", "m.chen@example.com", "817-555-0203", "456 Elm Ave, Fort Worth, TX 76102", "Linda Chen, 817-555-0204", "Yes", "Tuesday Afternoon - Library", "Friend/Family"],
    ["Rodriguez", "Emily", "emily.r@domain.com", "817-555-0305", "789 Pine Rd, Arlington, TX 76011", "Carlos Rodriguez, 817-555-0306", "Yes", "Wednesday Evening", "Church"],
    ["Thompson", "David", "david.t@mail.com", "817-555-0407", "321 Maple Ln, Grand Prairie, TX 75050", "Susan Thompson, 817-555-0408", "Yes", "Thursday Afternoon", "HOPE Website"],
    ["Williams", "Jessica", "jwilliams@company.com", "817-555-0509", "654 Cedar Ct, Mansfield, TX 76063", "Robert Williams, 817-555-0510", "Yes", "Saturday Morning", "Volunteer Match"],
    ["Garcia", "Antonio", "a.garcia@inbox.com", "817-555-0611", "987 Birch Dr, Arlington, TX 76012", "Maria Garcia, 817-555-0612", "Yes", "Monday Evening - UTA", "School/University"],
    ["Patel", "Priya", "priya.p@webmail.com", "817-555-0713", "147 Willow Way, Fort Worth, TX 76104", "Raj Patel, 817-555-0714", "Yes", "Tuesday Afternoon", "Employer"],
    ["Brown", "Christopher", "c.brown@email.net", "817-555-0815", "258 Spruce St, Arlington, TX 76013", "Amanda Brown, 817-555-0816", "Yes", "Wednesday Evening", "Other Volunteer"],
]

# Demo data - second batch of sample volunteers (Batch 2)
DEMO_DATA_BATCH2 = [
    ["Martinez", "Isabella", "isabella.m@email.com", "817-555-1001", "789 Sunset Blvd, Arlington, TX 76017", "Carlos Martinez, 817-555-1002", "Yes", "Monday Evening - UTA", "Instagram"],
    ["Kim", "Daniel", "d.kim@techcorp.com", "817-555-1103", "456 Innovation Dr, Fort Worth, TX 76109", "Jennifer Kim, 817-555-1104", "Yes", "Tuesday Afternoon - Library", "LinkedIn"],
    ["Washington", "Jasmine", "jasmine.w@gmail.com", "817-555-1205", "321 Heritage Lane, Arlington, TX 76014", "Marcus Washington, 817-555-1206", "Yes", "Wednesday Evening", "Church Announcement"],
    ["O'Brien", "Patrick", "pobrien@university.edu", "817-555-1307", "654 Campus Way, Arlington, TX 76019", "Shannon O'Brien, 817-555-1308", "Yes", "Thursday Afternoon", "Professor Recommendation"],
    ["Nguyen", "Lisa", "lisa.nguyen@company.org", "817-555-1409", "987 Commerce St, Grand Prairie, TX 75051", "Tuan Nguyen, 817-555-1410", "Yes", "Saturday Morning", "Community Event"],
    ["Foster", "Marcus", "m.foster@startup.io", "817-555-1511", "147 Tech Park Dr, Arlington, TX 76006", "Angela Foster, 817-555-1512", "Yes", "Monday Evening - UTA", "Twitter/X"],
    ["Ramirez", "Sofia", "sofia.ramirez@mail.com", "817-555-1613", "258 Garden View, Mansfield, TX 76063", "Miguel Ramirez, 817-555-1614", "Yes", "Tuesday Afternoon", "Friend Referral"],
    ["Anderson", "James", "j.anderson@corp.net", "817-555-1715", "369 Executive Blvd, Arlington, TX 76011", "Catherine Anderson, 817-555-1716", "Yes", "Wednesday Evening", "Company Volunteer Program"],
    ["Lee", "Hannah", "hannah.lee@outlook.com", "817-555-1817", "741 Oak Ridge Rd, Fort Worth, TX 76107", "David Lee, 817-555-1818", "Yes", "Thursday Afternoon", "Facebook"],
    ["Cooper", "William", "wcooper@engineering.com", "817-555-1919", "852 Industrial Ave, Arlington, TX 76015", "Sarah Cooper, 817-555-1920", "Yes", "Saturday Morning", "HOPE Website"],
]

# Demo data - third batch V3 (3 students for quick demo)
DEMO_DATA_BATCH3 = [
    ["Taylor", "Emma", "emma.taylor@student.edu", "817-555-2001", "123 College Ave, Arlington, TX 76013", "Robert Taylor, 817-555-2002", "Yes", "Monday Evening - UTA", "Campus Flyer"],
    ["Jackson", "Noah", "noah.j@university.edu", "817-555-2103", "456 Dorm Circle, Fort Worth, TX 76102", "Michelle Jackson, 817-555-2104", "Yes", "Wednesday Evening", "Professor"],
    ["White", "Olivia", "olivia.white@mail.com", "817-555-2205", "789 Student Housing Dr, Arlington, TX 76010", "Thomas White, 817-555-2206", "Yes", "Saturday Morning", "Friend Recommendation"],
]


def set_cell_shading(cell, color):
    """Set background color for a Word table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_csv_template(output_dir):
    """Create empty CSV template"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_EMPTY.csv')
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(HEADERS)
    return filepath


def create_csv_demo(output_dir):
    """Create CSV with demo data"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_DEMO.csv')
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(HEADERS)
        writer.writerows(DEMO_DATA)
    return filepath


def create_docx_template(output_dir):
    """Create empty DOCX template - landscape, readable format"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_EMPTY.docx')
    
    doc = Document()
    
    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Title
    title = doc.add_heading('HOPE Tutoring - Volunteer Data Template', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Fill in volunteer information below. Upload to the Messaging Tool to generate emails.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    link = doc.add_paragraph('Website: https://www.hopetutoring.org/')
    link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Create table - 11 rows (1 header + 10 data rows)
    table = doc.add_table(rows=11, cols=len(SHORT_HEADERS))
    table.style = 'Table Grid'
    
    # Column widths optimized for landscape (total ~10 inches)
    col_widths = [Inches(0.85), Inches(0.85), Inches(1.4), Inches(0.95), Inches(1.5), Inches(1.5), Inches(0.45), Inches(1.1), Inches(1.5)]
    
    # Set column widths
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
    
    # Style header row
    for i, header in enumerate(SHORT_HEADERS):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '1a8b8d')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Set data row heights
    for row_idx in range(1, 11):
        table.rows[row_idx].height = Inches(0.35)
    
    doc.add_paragraph()
    
    # Add note
    note = doc.add_paragraph()
    note.add_run('Note: ').bold = True
    note.add_run('Volunteers can type their own answer for "How did you hear about us?" (Social Media, Friend, Church, School, Website, etc.)')
    
    doc.save(filepath)
    return filepath


def create_docx_demo(output_dir):
    """Create DOCX with demo data - landscape format"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_DEMO.docx')
    
    doc = Document()
    
    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Title
    title = doc.add_heading('HOPE Tutoring - Volunteer Data (Demo)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Sample data for testing the Adaptive Messaging Script')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=len(DEMO_DATA) + 1, cols=len(SHORT_HEADERS))
    table.style = 'Table Grid'
    
    col_widths = [Inches(0.85), Inches(0.85), Inches(1.4), Inches(0.95), Inches(1.5), Inches(1.5), Inches(0.45), Inches(1.1), Inches(1.5)]
    
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
    
    # Header row
    for i, header in enumerate(SHORT_HEADERS):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '1a8b8d')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for row_idx, row_data in enumerate(DEMO_DATA):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    doc.save(filepath)
    return filepath


def create_pdf_template(output_dir):
    """Create empty PDF template - two-page layout for readability"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_EMPTY.pdf')
    
    doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, alignment=1, spaceAfter=10, textColor=colors.HexColor('#1a8b8d'))
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=11, alignment=1, spaceAfter=15)
    header_style = ParagraphStyle('Header', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=1)
    note_style = ParagraphStyle('Note', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#666666'))
    
    elements = []
    
    # Page 1: Personal Info
    elements.append(Paragraph('HOPE Tutoring - Volunteer Data Template', title_style))
    elements.append(Paragraph('Fill in below. Upload to Messaging Tool. | www.hopetutoring.org', subtitle_style))
    elements.append(Paragraph('<b>Part 1: Personal Information</b>', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    p1_headers = ['Last Name', 'First Name', 'Email Address', 'Phone Number', 'Address (City, Zip)']
    p1_data = [[Paragraph(h, header_style) for h in p1_headers]]
    for _ in range(12):
        p1_data.append(['', '', '', '', ''])
    
    t1 = Table(p1_data, colWidths=[1.0*inch, 1.0*inch, 1.8*inch, 1.1*inch, 2.4*inch], repeatRows=1)
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a8b8d')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
        ('ROWHEIGHT', (0, 1), (-1, -1), 30),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(t1)
    elements.append(PageBreak())
    
    # Page 2: Additional Info
    elements.append(Paragraph('HOPE Tutoring - Volunteer Data Template', title_style))
    elements.append(Paragraph('<b>Part 2: Additional Information</b>', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    p2_headers = ['Emergency Contact (Name & Phone)', '18+?', 'Site/Session', 'How did you hear about us?']
    p2_data = [[Paragraph(h, header_style) for h in p2_headers]]
    for _ in range(12):
        p2_data.append(['', '', '', ''])
    
    t2 = Table(p2_data, colWidths=[2.3*inch, 0.6*inch, 1.8*inch, 2.6*inch], repeatRows=1)
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a8b8d')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
        ('ROWHEIGHT', (0, 1), (-1, -1), 30),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(t2)
    elements.append(Spacer(1, 20))
    elements.append(Paragraph('<b>Note:</b> Type your own answer for "How did you hear about us?" (Social Media, Friend, Church, School, Website, etc.)', note_style))
    
    doc.build(elements)
    return filepath


def create_pdf_demo(output_dir):
    """Create PDF with demo data - two-page layout"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_DEMO.pdf')
    
    doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, alignment=1, spaceAfter=10, textColor=colors.HexColor('#1a8b8d'))
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=11, alignment=1, spaceAfter=15)
    header_style = ParagraphStyle('Header', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=1)
    cell_style = ParagraphStyle('Cell', parent=styles['Normal'], fontSize=8, leading=10)
    
    elements = []
    
    # Page 1
    elements.append(Paragraph('HOPE Tutoring - Volunteer Data (Demo)', title_style))
    elements.append(Paragraph('Sample data for testing', subtitle_style))
    elements.append(Paragraph('<b>Part 1: Personal Information</b>', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    p1_headers = ['Last Name', 'First Name', 'Email', 'Phone', 'Address']
    p1_data = [[Paragraph(h, header_style) for h in p1_headers]]
    for row in DEMO_DATA:
        p1_data.append([Paragraph(row[i], cell_style) for i in range(5)])
    
    t1 = Table(p1_data, colWidths=[0.9*inch, 0.9*inch, 1.7*inch, 1.0*inch, 2.8*inch], repeatRows=1)
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a8b8d')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
        ('ROWHEIGHT', (0, 1), (-1, -1), 26),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUND', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5fafa')]),
    ]))
    elements.append(t1)
    elements.append(PageBreak())
    
    # Page 2
    elements.append(Paragraph('HOPE Tutoring - Volunteer Data (Demo)', title_style))
    elements.append(Paragraph('<b>Part 2: Additional Information</b>', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    p2_headers = ['Emergency Contact', '18+?', 'Site/Session', 'How heard?']
    p2_data = [[Paragraph(h, header_style) for h in p2_headers]]
    for row in DEMO_DATA:
        p2_data.append([Paragraph(row[i], cell_style) for i in range(5, 9)])
    
    t2 = Table(p2_data, colWidths=[2.3*inch, 0.5*inch, 1.8*inch, 2.7*inch], repeatRows=1)
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a8b8d')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
        ('ROWHEIGHT', (0, 1), (-1, -1), 26),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUND', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5fafa')]),
    ]))
    elements.append(t2)
    
    doc.build(elements)
    return filepath


def create_csv_demo_batch2(output_dir):
    """Create CSV with batch 2 demo data"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_DEMO_BATCH2.csv')
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(HEADERS)
        writer.writerows(DEMO_DATA_BATCH2)
    return filepath


def create_docx_demo_batch2(output_dir):
    """Create DOCX with batch 2 demo data - landscape format"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_DEMO_BATCH2.docx')
    
    doc = Document()
    
    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Title
    title = doc.add_heading('HOPE Tutoring - Volunteer Data (Demo Batch 2)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Second batch of sample data for testing the HSRM')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=len(DEMO_DATA_BATCH2) + 1, cols=len(SHORT_HEADERS))
    table.style = 'Table Grid'
    
    col_widths = [Inches(0.85), Inches(0.85), Inches(1.4), Inches(0.95), Inches(1.5), Inches(1.5), Inches(0.45), Inches(1.1), Inches(1.5)]
    
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
    
    # Header row
    for i, header in enumerate(SHORT_HEADERS):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '5AAFE0')  # HOPE blue
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for row_idx, row_data in enumerate(DEMO_DATA_BATCH2):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    doc.save(filepath)
    return filepath


def create_pdf_demo_batch2(output_dir):
    """Create PDF with batch 2 demo data - two-page layout"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_DEMO_BATCH2.pdf')
    
    doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, alignment=1, spaceAfter=10, textColor=colors.HexColor('#5AAFE0'))
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=11, alignment=1, spaceAfter=15)
    header_style = ParagraphStyle('Header', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=1)
    cell_style = ParagraphStyle('Cell', parent=styles['Normal'], fontSize=8, leading=10)
    
    elements = []
    
    # Page 1
    elements.append(Paragraph('HOPE Tutoring - Volunteer Data (Demo Batch 2)', title_style))
    elements.append(Paragraph('Second batch for testing', subtitle_style))
    elements.append(Paragraph('<b>Part 1: Personal Information</b>', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    p1_headers = ['Last Name', 'First Name', 'Email', 'Phone', 'Address']
    p1_data = [[Paragraph(h, header_style) for h in p1_headers]]
    for row in DEMO_DATA_BATCH2:
        p1_data.append([Paragraph(row[i], cell_style) for i in range(5)])
    
    t1 = Table(p1_data, colWidths=[0.9*inch, 0.9*inch, 1.7*inch, 1.0*inch, 2.8*inch], repeatRows=1)
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#5AAFE0')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
        ('ROWHEIGHT', (0, 1), (-1, -1), 26),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUND', (0, 1), (-1, -1), [colors.white, colors.HexColor('#E8F4FB')]),
    ]))
    elements.append(t1)
    elements.append(PageBreak())
    
    # Page 2
    elements.append(Paragraph('HOPE Tutoring - Volunteer Data (Demo Batch 2)', title_style))
    elements.append(Paragraph('<b>Part 2: Additional Information</b>', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    p2_headers = ['Emergency Contact', '18+?', 'Site/Session', 'How heard?']
    p2_data = [[Paragraph(h, header_style) for h in p2_headers]]
    for row in DEMO_DATA_BATCH2:
        p2_data.append([Paragraph(row[i], cell_style) for i in range(5, 9)])
    
    t2 = Table(p2_data, colWidths=[2.3*inch, 0.5*inch, 1.8*inch, 2.7*inch], repeatRows=1)
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#5AAFE0')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
        ('ROWHEIGHT', (0, 1), (-1, -1), 26),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUND', (0, 1), (-1, -1), [colors.white, colors.HexColor('#E8F4FB')]),
    ]))
    elements.append(t2)
    
    doc.build(elements)
    return filepath


def create_csv_demo_batch3(output_dir):
    """Create CSV with batch 3 demo data (3 students)"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_DEMO_V3.csv')
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(HEADERS)
        writer.writerows(DEMO_DATA_BATCH3)
    return filepath


def create_docx_demo_batch3(output_dir):
    """Create DOCX with batch 3 demo data - 3 students"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_DEMO_V3.docx')
    
    doc = Document()
    
    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Title
    title = doc.add_heading('HOPE Tutoring - Volunteer Data (Demo V3)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Quick demo batch - 3 student volunteers')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=len(DEMO_DATA_BATCH3) + 1, cols=len(SHORT_HEADERS))
    table.style = 'Table Grid'
    
    col_widths = [Inches(0.85), Inches(0.85), Inches(1.4), Inches(0.95), Inches(1.5), Inches(1.5), Inches(0.45), Inches(1.1), Inches(1.5)]
    
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
    
    # Header row - purple color for V3
    for i, header in enumerate(SHORT_HEADERS):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '9B59B6')  # Purple
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for row_idx, row_data in enumerate(DEMO_DATA_BATCH3):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    doc.save(filepath)
    return filepath


def create_pdf_demo_batch3(output_dir):
    """Create PDF with batch 3 demo data - 3 students, two-page layout (same as Batch 1 & 2)"""
    filepath = os.path.join(output_dir, 'Volunteer_Template_DEMO_V3.pdf')
    
    doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, alignment=1, spaceAfter=10, textColor=colors.HexColor('#9B59B6'))
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=11, alignment=1, spaceAfter=15)
    header_style = ParagraphStyle('Header', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=1)
    cell_style = ParagraphStyle('Cell', parent=styles['Normal'], fontSize=8, leading=10)
    
    elements = []
    
    # Page 1 - Personal Information
    elements.append(Paragraph('HOPE Tutoring - Volunteer Data (Demo V3)', title_style))
    elements.append(Paragraph('Quick demo batch - 3 student volunteers', subtitle_style))
    elements.append(Paragraph('<b>Part 1: Personal Information</b>', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    p1_headers = ['Last Name', 'First Name', 'Email', 'Phone', 'Address']
    p1_data = [[Paragraph(h, header_style) for h in p1_headers]]
    for row in DEMO_DATA_BATCH3:
        p1_data.append([Paragraph(row[i], cell_style) for i in range(5)])
    
    t1 = Table(p1_data, colWidths=[0.9*inch, 0.9*inch, 1.7*inch, 1.0*inch, 2.8*inch], repeatRows=1)
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#9B59B6')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
        ('ROWHEIGHT', (0, 1), (-1, -1), 26),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUND', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5EEF8')]),
    ]))
    elements.append(t1)
    elements.append(PageBreak())
    
    # Page 2 - Additional Information
    elements.append(Paragraph('HOPE Tutoring - Volunteer Data (Demo V3)', title_style))
    elements.append(Paragraph('<b>Part 2: Additional Information</b>', styles['Heading2']))
    elements.append(Spacer(1, 10))
    
    p2_headers = ['Emergency Contact', '18+?', 'Site/Session', 'How heard?']
    p2_data = [[Paragraph(h, header_style) for h in p2_headers]]
    for row in DEMO_DATA_BATCH3:
        p2_data.append([Paragraph(row[i], cell_style) for i in range(5, 9)])
    
    t2 = Table(p2_data, colWidths=[2.3*inch, 0.5*inch, 1.8*inch, 2.7*inch], repeatRows=1)
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#9B59B6')),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
        ('ROWHEIGHT', (0, 1), (-1, -1), 26),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUND', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5EEF8')]),
    ]))
    elements.append(t2)
    elements.append(Spacer(1, 20))
    
    # Note
    note_style = ParagraphStyle('Note', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#666666'))
    elements.append(Paragraph('<b>Note:</b> This is a quick 3-student demo batch for testing batch tracking features.', note_style))
    
    doc.build(elements)
    return filepath


if __name__ == '__main__':
    # Allow custom output directory via command line
    output_dir = sys.argv[1] if len(sys.argv) > 1 else 'Admin Templates'
    
    # Create output directory
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"\n📁 Generating HOPE Tutoring Volunteer Templates to '{output_dir}'...\n")
    
    # Generate all templates (Batch 1)
    print("📋 Batch 1 - Empty & Original Demo:")
    print(f"   ✓ Created: {create_csv_template(output_dir)}")
    print(f"   ✓ Created: {create_csv_demo(output_dir)}")
    print(f"   ✓ Created: {create_docx_template(output_dir)}")
    print(f"   ✓ Created: {create_docx_demo(output_dir)}")
    print(f"   ✓ Created: {create_pdf_template(output_dir)}")
    print(f"   ✓ Created: {create_pdf_demo(output_dir)}")
    
    # Generate Batch 2 demo files
    print("\n📋 Batch 2 - Additional Demo Data:")
    print(f"   ✓ Created: {create_csv_demo_batch2(output_dir)}")
    print(f"   ✓ Created: {create_docx_demo_batch2(output_dir)}")
    print(f"   ✓ Created: {create_pdf_demo_batch2(output_dir)}")
    
    # Generate Batch 3 (V3) demo files
    print("\n📋 Batch 3 (V3) - Quick Demo (3 students):")
    print(f"   ✓ Created: {create_csv_demo_batch3(output_dir)}")
    print(f"   ✓ Created: {create_docx_demo_batch3(output_dir)}")
    print(f"   ✓ Created: {create_pdf_demo_batch3(output_dir)}")
    
    print("\n✅ All templates created successfully!")
    print(f"\n📂 Files in '{output_dir}':")
    print("   Empty Templates:")
    print("   • Volunteer_Template_EMPTY.csv  - Empty CSV")
    print("   • Volunteer_Template_EMPTY.docx - Empty Word (landscape)")
    print("   • Volunteer_Template_EMPTY.pdf  - Empty PDF (2-page)")
    print("\n   Demo Batch 1 (8 volunteers):")
    print("   • Volunteer_Template_DEMO.csv   - CSV with sample data")
    print("   • Volunteer_Template_DEMO.docx  - Word with data")
    print("   • Volunteer_Template_DEMO.pdf   - PDF with data")
    print("\n   Demo Batch 2 (10 different volunteers):")
    print("   • Volunteer_Template_DEMO_BATCH2.csv  - Second batch CSV")
    print("   • Volunteer_Template_DEMO_BATCH2.docx - Second batch Word")
    print("   • Volunteer_Template_DEMO_BATCH2.pdf  - Second batch PDF")
    print("\n   Demo V3 (3 students - quick demo):")
    print("   • Volunteer_Template_DEMO_V3.csv  - V3 CSV")
    print("   • Volunteer_Template_DEMO_V3.docx - V3 Word")
    print("   • Volunteer_Template_DEMO_V3.pdf  - V3 PDF")
