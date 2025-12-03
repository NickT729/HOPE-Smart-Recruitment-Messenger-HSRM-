"""
File Parser Module
Handles extraction of volunteer data from CSV, DOCX, and PDF files.
"""

import re
import pandas as pd
import pdfplumber
from docx import Document


class FileParser:
    """Parse various file formats to extract volunteer contact information."""
    
    # Common column name mappings (includes HOPE form fields)
    FIRST_NAME_COLUMNS = ['first_name', 'firstname', 'first', 'fname']
    LAST_NAME_COLUMNS = ['last_name', 'lastname', 'last', 'lname', 'surname']
    NAME_COLUMNS = ['name', 'full_name', 'fullname', 'volunteer_name', 'volunteer']
    EMAIL_COLUMNS = ['email', 'e-mail', 'email_address', 'emailaddress', 'mail', 'contact_email']
    PHONE_COLUMNS = ['phone', 'telephone', 'phone_number', 'phonenumber', 'mobile', 'cell', 'contact']
    INTEREST_COLUMNS = ['interest', 'interests', 'area_of_interest', 'subject', 'subjects', 'skills', 'expertise', 
                        'assigned_site/session', 'assigned_site', 'site', 'session']
    LOCATION_COLUMNS = ['location', 'city', 'area', 'neighborhood', 'address', 'zip', 'zipcode',
                        'street_address_with_city_&_zip_code', 'street_address', 'street']
    REFERRAL_COLUMNS = ['how_did_you_hear_about_hope_tutoring?', 'how_did_you_hear', 'referral', 'source', 'heard_about']
    
    # Email regex pattern
    EMAIL_PATTERN = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
    
    # Phone regex pattern
    PHONE_PATTERN = re.compile(r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}')
    
    def parse(self, filepath: str) -> list:
        """
        Parse a file and extract volunteer data.
        
        Args:
            filepath: Path to the file to parse
            
        Returns:
            List of dictionaries containing volunteer data
        """
        extension = filepath.rsplit('.', 1)[1].lower()
        
        if extension == 'csv':
            return self._parse_csv(filepath)
        elif extension == 'docx':
            return self._parse_docx(filepath)
        elif extension == 'pdf':
            return self._parse_pdf(filepath)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _parse_csv(self, filepath: str) -> list:
        """Parse CSV file and extract volunteer data."""
        try:
            df = pd.read_csv(filepath)
            return self._extract_from_dataframe(df)
        except Exception as e:
            raise ValueError(f"Error parsing CSV: {str(e)}")
    
    def _parse_docx(self, filepath: str) -> list:
        """Parse Word document and extract volunteer data."""
        try:
            doc = Document(filepath)
            volunteers = []
            
            # First, try to find tables
            for table in doc.tables:
                df = self._table_to_dataframe(table)
                if df is not None and not df.empty:
                    volunteers.extend(self._extract_from_dataframe(df))
            
            # If no tables found, try to extract from text
            if not volunteers:
                text = '\n'.join([para.text for para in doc.paragraphs])
                volunteers = self._extract_from_text(text)
            
            return volunteers
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
    
    def _parse_pdf(self, filepath: str) -> list:
        """Parse PDF file and extract volunteer data."""
        try:
            volunteers = []
            
            with pdfplumber.open(filepath) as pdf:
                # First, try to extract tables
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            df = pd.DataFrame(table[1:], columns=table[0] if table[0] else None)
                            extracted = self._extract_from_dataframe(df)
                            volunteers.extend(extracted)
                
                # If no tables found, try to extract from text
                if not volunteers:
                    text = '\n'.join([page.extract_text() or '' for page in pdf.pages])
                    volunteers = self._extract_from_text(text)
            
            return volunteers
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    def _table_to_dataframe(self, table) -> pd.DataFrame:
        """Convert a Word table to a pandas DataFrame."""
        data = []
        headers = None
        
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            if i == 0:
                headers = row_data
            else:
                data.append(row_data)
        
        if headers and data:
            return pd.DataFrame(data, columns=headers)
        return pd.DataFrame()
    
    def _extract_from_dataframe(self, df: pd.DataFrame) -> list:
        """Extract volunteer data from a DataFrame."""
        if df.empty:
            return []
        
        # Normalize column names
        df.columns = [str(col).lower().strip().replace(' ', '_') for col in df.columns]
        
        # Find relevant columns
        first_name_col = self._find_column(df.columns, self.FIRST_NAME_COLUMNS)
        last_name_col = self._find_column(df.columns, self.LAST_NAME_COLUMNS)
        name_col = self._find_column(df.columns, self.NAME_COLUMNS)
        email_col = self._find_column(df.columns, self.EMAIL_COLUMNS)
        phone_col = self._find_column(df.columns, self.PHONE_COLUMNS)
        interest_col = self._find_column(df.columns, self.INTEREST_COLUMNS)
        location_col = self._find_column(df.columns, self.LOCATION_COLUMNS)
        referral_col = self._find_column(df.columns, self.REFERRAL_COLUMNS)
        
        volunteers = []
        
        for _, row in df.iterrows():
            # Try to get email - this is required
            email = None
            if email_col:
                email = str(row[email_col]).strip()
            else:
                # Try to find email in any column
                for col in df.columns:
                    val = str(row[col])
                    match = self.EMAIL_PATTERN.search(val)
                    if match:
                        email = match.group()
                        break
            
            # Skip rows without email
            if not email or email == 'nan' or '@' not in email:
                continue
            
            # Get name - handle both separate first/last name columns and combined name
            first_name = ''
            last_name = ''
            name = ''
            
            # Check for separate first/last name columns (HOPE form format)
            if first_name_col:
                first_name = str(row[first_name_col]).strip()
                if first_name == 'nan':
                    first_name = ''
            
            if last_name_col:
                last_name = str(row[last_name_col]).strip()
                if last_name == 'nan':
                    last_name = ''
            
            # If we have first/last name, combine them
            if first_name or last_name:
                name = f"{first_name} {last_name}".strip()
            # Otherwise try combined name column
            elif name_col:
                name = str(row[name_col]).strip()
                if name == 'nan':
                    name = ''
                # Extract first name from combined name
                if name:
                    first_name = name.split()[0]
            
            # If still no name, try first column that's not email
            if not name:
                for col in df.columns:
                    val = str(row[col]).strip()
                    if val and val != 'nan' and '@' not in val and not self.PHONE_PATTERN.match(val):
                        name = val
                        first_name = val.split()[0] if val else ''
                        break
            
            volunteer = {
                'name': name,
                'email': email,
                'first_name': first_name if first_name else (name.split()[0] if name else ''),
                'last_name': last_name,
                'phone': '',
                'interests': '',
                'location': '',
                'referral': ''
            }
            
            # Get optional fields
            if phone_col and str(row[phone_col]) != 'nan':
                volunteer['phone'] = str(row[phone_col]).strip()
            
            if interest_col and str(row[interest_col]) != 'nan':
                volunteer['interests'] = str(row[interest_col]).strip()
            
            if location_col and str(row[location_col]) != 'nan':
                volunteer['location'] = str(row[location_col]).strip()
            
            if referral_col and str(row[referral_col]) != 'nan':
                volunteer['referral'] = str(row[referral_col]).strip()
            
            volunteers.append(volunteer)
        
        return volunteers
    
    def _extract_from_text(self, text: str) -> list:
        """Extract volunteer data from unstructured text."""
        volunteers = []
        
        # Find all emails in the text
        emails = self.EMAIL_PATTERN.findall(text)
        
        # For each email, try to find associated name
        lines = text.split('\n')
        
        for email in set(emails):  # Use set to avoid duplicates
            name = ''
            
            # Look for name near the email
            for i, line in enumerate(lines):
                if email in line:
                    # Check current line for name before email
                    parts = line.split(email)[0].strip()
                    if parts:
                        # Clean up the name
                        name = re.sub(r'[^\w\s]', '', parts).strip()
                    
                    # If no name found, check previous line
                    if not name and i > 0:
                        prev_line = lines[i-1].strip()
                        if prev_line and '@' not in prev_line:
                            name = re.sub(r'[^\w\s]', '', prev_line).strip()
                    break
            
            volunteer = {
                'name': name,
                'email': email,
                'first_name': name.split()[0] if name else '',
                'last_name': '',
                'phone': '',
                'interests': '',
                'location': '',
                'referral': ''
            }
            
            volunteers.append(volunteer)
        
        return volunteers
    
    def _find_column(self, columns, possible_names: list) -> str:
        """Find a column name from a list of possible names."""
        columns_lower = [str(col).lower() for col in columns]
        
        for name in possible_names:
            if name in columns_lower:
                idx = columns_lower.index(name)
                return list(columns)[idx]
            # Check for partial match
            for i, col in enumerate(columns_lower):
                if name in col:
                    return list(columns)[i]
        
        return None

